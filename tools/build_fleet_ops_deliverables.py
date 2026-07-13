from pathlib import Path
from datetime import date
from xml.sax.saxutils import escape

from PIL import Image, ImageDraw, ImageFont
from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
OUT = ROOT / "deliverables"
QA = OUT / "qa"
OUT.mkdir(exist_ok=True)
QA.mkdir(exist_ok=True)

DOCX_PATH = OUT / "Fleet_Ops_Copilot_Project_Report.docx"
DRAWIO_PATH = OUT / "Fleet_Ops_Copilot_Design.drawio"
ARCH_PNG = QA / "fleet_ops_architecture.png"

NAVY = "0B2545"
BLUE = "2E74B5"
DARK_BLUE = "1F4D78"
SKY = "E8F2FA"
PALE = "F4F6F9"
GOLD = "D9A441"
MUTED = "5E6B78"
WHITE = "FFFFFF"
BLACK = "222222"
GREEN = "2F855A"
RED = "9B1C1C"


def font(name="Calibri", size=22, bold=False):
    candidates = [
        Path("C:/Windows/Fonts/calibrib.ttf" if bold else "C:/Windows/Fonts/calibri.ttf"),
        Path("C:/Windows/Fonts/arialbd.ttf" if bold else "C:/Windows/Fonts/arial.ttf"),
    ]
    for candidate in candidates:
        if candidate.exists():
            return ImageFont.truetype(str(candidate), size)
    return ImageFont.load_default()


def draw_architecture_png():
    img = Image.new("RGB", (1800, 1080), "#FFFFFF")
    d = ImageDraw.Draw(img)
    title = font(size=42, bold=True)
    label = font(size=25, bold=True)
    body = font(size=20)
    small = font(size=17)
    d.text((70, 45), "Fleet Ops Copilot - Reference Architecture", fill="#0B2545", font=title)
    d.text((70, 102), "Conversational access to live transportation operations, schedules, and guided resolution", fill="#5E6B78", font=body)

    def box(x, y, w, h, heading, lines, fill, outline="#2E74B5"):
        d.rounded_rectangle((x, y, x+w, y+h), radius=18, fill=fill, outline=outline, width=4)
        d.text((x+22, y+18), heading, fill="#0B2545", font=label)
        yy = y + 58
        for line in lines:
            d.text((x+24, yy), line, fill="#263746", font=small)
            yy += 27

    def arrow(x1, y1, x2, y2, color="#2E74B5"):
        d.line((x1, y1, x2, y2), fill=color, width=5)
        import math
        a = math.atan2(y2-y1, x2-x1)
        size = 14
        pts = [(x2, y2), (x2-size*math.cos(a-0.5), y2-size*math.sin(a-0.5)), (x2-size*math.cos(a+0.5), y2-size*math.sin(a+0.5))]
        d.polygon(pts, fill=color)

    box(70, 190, 300, 190, "Users & Channels", ["Dispatcher / planner", "Driver / field team", "Operations manager", "Web, mobile, chat"], "#E8F2FA")
    box(485, 170, 500, 230, "Copilot Experience", ["Conversation UI & session context", "Identity, role and consent", "Answer with evidence and timestamps", "Human approval for operational actions"], "#F4F6F9")
    box(1100, 170, 620, 230, "Agent Orchestration", ["Intent classification & planning", "Tool selection and guarded execution", "Retrieval, grounding and response synthesis", "Policy checks, confidence and escalation"], "#FFF8E8", outline="#D9A441")

    box(120, 520, 420, 230, "Operational Tools / APIs", ["Schedule & route service", "Vehicle / driver availability", "Incident & ticket management", "Traffic, weather and map feeds"], "#EDF8F1", outline="#2F855A")
    box(690, 520, 420, 230, "Data & Knowledge", ["Real-time operational database", "Policies, SOPs and route knowledge", "Vector search / retrieval index", "Conversation and audit records"], "#E8F2FA")
    box(1260, 520, 410, 230, "Platform Controls", ["API gateway and access control", "Observability and cost controls", "Encryption and secrets", "Audit, retention and compliance"], "#FDEEEE", outline="#9B1C1C")

    arrow(370, 285, 485, 285)
    arrow(985, 285, 1100, 285)
    arrow(1310, 400, 390, 520)
    arrow(1390, 400, 900, 520)
    arrow(1490, 400, 1465, 520)
    arrow(540, 650, 690, 650)
    arrow(1260, 650, 1110, 650)

    d.rounded_rectangle((70, 850, 1650, 1050), radius=18, fill="#0B2545")
    d.text((100, 875), "Response contract", fill="#FFFFFF", font=label)
    contract = [
        "1. State the answer and operational impact clearly.",
        "2. Show source system, effective time, confidence, and assumptions.",
        "3. Recommend the next safe action; request approval before making changes.",
        "4. Escalate when data is stale, tools fail, or the decision exceeds policy.",
    ]
    yy = 918
    for line in contract:
        d.text((115, yy), line, fill="#FFFFFF", font=small)
        yy += 29
    img.save(ARCH_PNG, quality=95)


def drawio_cell(cid, value, style, x, y, w, h, parent="1", vertex=True):
    value = escape(value)
    attr = 'vertex="1"' if vertex else 'edge="1"'
    return f'<mxCell id="{cid}" value="{value}" style="{style}" {attr} parent="{parent}"><mxGeometry x="{x}" y="{y}" width="{w}" height="{h}" as="geometry"/></mxCell>'


def drawio_edge(cid, source, target, value=""):
    style = "edgeStyle=orthogonalEdgeStyle;rounded=1;orthogonalLoop=1;jettySize=auto;html=1;endArrow=block;endFill=1;strokeWidth=2;strokeColor=#2E74B5;"
    return f'<mxCell id="{cid}" value="{escape(value)}" style="{style}" edge="1" parent="1" source="{source}" target="{target}"><mxGeometry relative="1" as="geometry"/></mxCell>'


def build_drawio():
    title_style = "text;html=1;strokeColor=none;fillColor=none;align=left;verticalAlign=middle;fontSize=26;fontStyle=1;fontColor=#0B2545;"
    box_blue = "rounded=1;whiteSpace=wrap;html=1;fillColor=#E8F2FA;strokeColor=#2E74B5;strokeWidth=2;fontSize=15;fontColor=#0B2545;spacing=12;"
    box_gray = "rounded=1;whiteSpace=wrap;html=1;fillColor=#F4F6F9;strokeColor=#2E74B5;strokeWidth=2;fontSize=15;fontColor=#0B2545;spacing=12;"
    box_gold = "rounded=1;whiteSpace=wrap;html=1;fillColor=#FFF8E8;strokeColor=#D9A441;strokeWidth=2;fontSize=15;fontColor=#0B2545;spacing=12;"
    box_green = "rounded=1;whiteSpace=wrap;html=1;fillColor=#EDF8F1;strokeColor=#2F855A;strokeWidth=2;fontSize=15;fontColor=#0B2545;spacing=12;"
    box_red = "rounded=1;whiteSpace=wrap;html=1;fillColor=#FDEEEE;strokeColor=#9B1C1C;strokeWidth=2;fontSize=15;fontColor=#0B2545;spacing=12;"
    cells = [
        '<mxCell id="0"/>', '<mxCell id="1" parent="0"/>',
        drawio_cell("t", "Fleet Ops Copilot - System Architecture", title_style, 40, 20, 800, 45),
        drawio_cell("users", "<b>Users &amp; Channels</b><br>Dispatcher / Planner<br>Driver / Field Team<br>Operations Manager<br>Web / Mobile / Chat", box_blue, 40, 120, 230, 180),
        drawio_cell("ux", "<b>Copilot Experience</b><br>Conversation UI<br>Session context<br>Identity and role<br>Evidence and timestamps", box_gray, 360, 100, 300, 220),
        drawio_cell("agent", "<b>Agent Orchestration</b><br>Intent and planning<br>Tool selection<br>Grounded synthesis<br>Policy and confidence<br>Human approval gate", box_gold, 750, 90, 340, 240),
        drawio_cell("tools", "<b>Operational Tools / APIs</b><br>Schedule and route<br>Vehicle and driver<br>Incident and ticket<br>Traffic, weather, maps", box_green, 110, 450, 300, 220),
        drawio_cell("data", "<b>Data &amp; Knowledge</b><br>Operational database<br>Policies and SOPs<br>Retrieval index<br>Conversation and audit", box_blue, 500, 450, 300, 220),
        drawio_cell("controls", "<b>Platform Controls</b><br>API gateway / RBAC<br>Observability<br>Encryption and secrets<br>Audit and retention", box_red, 890, 450, 300, 220),
        drawio_edge("e1", "users", "ux"), drawio_edge("e2", "ux", "agent"),
        drawio_edge("e3", "agent", "tools", "tool calls"), drawio_edge("e4", "agent", "data", "retrieve / store"),
        drawio_edge("e5", "agent", "controls", "enforce / observe"), drawio_edge("e6", "tools", "data", "live state"),
    ]
    arch = '<diagram id="architecture" name="System Architecture"><mxGraphModel dx="1422" dy="794" grid="1" gridSize="10" guides="1" tooltips="1" connect="1" arrows="1" fold="1" page="1" pageScale="1" pageWidth="1300" pageHeight="800" math="0" shadow="0"><root>' + ''.join(cells) + '</root></mxGraphModel></diagram>'

    seq_style = "rounded=1;whiteSpace=wrap;html=1;fillColor=#E8F2FA;strokeColor=#2E74B5;strokeWidth=2;fontSize=15;fontColor=#0B2545;spacing=10;"
    gate_style = "rhombus;whiteSpace=wrap;html=1;fillColor=#FFF8E8;strokeColor=#D9A441;strokeWidth=2;fontSize=14;fontColor=#0B2545;"
    cells2 = ['<mxCell id="0"/>', '<mxCell id="1" parent="0"/>', drawio_cell("t2", "Fleet Ops Copilot - Agent Workflow", title_style, 40, 20, 700, 45)]
    steps = [
        ("s1", "1. User request", 60, 120), ("s2", "2. Authenticate &amp;<br>load role/context", 300, 120),
        ("s3", "3. Detect intent &amp;<br>required freshness", 560, 120), ("s4", "4. Plan and choose<br>approved tools", 820, 120),
        ("s5", "5. Query live systems<br>and retrieve SOPs", 820, 350), ("s6", "6. Validate sources,<br>time and conflicts", 560, 350),
        ("gate", "Operational<br>change?", 320, 330), ("s7", "7. Explain answer &amp;<br>recommended action", 60, 350),
        ("s8", "8. Request approval,<br>execute, audit", 300, 560), ("s9", "9. Confirm outcome<br>or escalate", 60, 560),
    ]
    for cid, val, x, y in steps:
        cells2.append(drawio_cell(cid, val, gate_style if cid == "gate" else seq_style, x, y, 180, 100))
    for i, (a,b,label) in enumerate([
        ("s1","s2",""),("s2","s3",""),("s3","s4",""),("s4","s5",""),("s5","s6",""),("s6","gate",""),
        ("gate","s7","No"),("gate","s8","Yes"),("s8","s9",""),("s7","s9","")
    ], start=20): cells2.append(drawio_edge(f"e{i}", a, b, label))
    workflow = '<diagram id="workflow" name="Agent Workflow"><mxGraphModel dx="1422" dy="794" grid="1" gridSize="10" guides="1" tooltips="1" connect="1" arrows="1" fold="1" page="1" pageScale="1" pageWidth="1100" pageHeight="750" math="0" shadow="0"><root>' + ''.join(cells2) + '</root></mxGraphModel></diagram>'
    xml = '<?xml version="1.0" encoding="UTF-8"?><mxfile host="app.diagrams.net" modified="2026-07-13T00:00:00.000Z" agent="Codex" version="24.7.17" type="device">' + arch + workflow + '</mxfile>'
    DRAWIO_PATH.write_text(xml, encoding="utf-8")


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=80, start=120, bottom=80, end=120):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcMar = tcPr.first_child_found_in("w:tcMar")
    if tcMar is None:
        tcMar = OxmlElement("w:tcMar")
        tcPr.append(tcMar)
    for m, v in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tcMar.find(qn(f"w:{m}"))
        if node is None:
            node = OxmlElement(f"w:{m}")
            tcMar.append(node)
        node.set(qn("w:w"), str(v)); node.set(qn("w:type"), "dxa")


def set_table_geometry(table, widths_dxa):
    table.autofit = False
    tblPr = table._tbl.tblPr
    tblW = tblPr.find(qn("w:tblW"))
    if tblW is None:
        tblW = OxmlElement("w:tblW"); tblPr.append(tblW)
    tblW.set(qn("w:w"), str(sum(widths_dxa))); tblW.set(qn("w:type"), "dxa")
    tblInd = tblPr.find(qn("w:tblInd"))
    if tblInd is None:
        tblInd = OxmlElement("w:tblInd"); tblPr.append(tblInd)
    tblInd.set(qn("w:w"), "120"); tblInd.set(qn("w:type"), "dxa")
    grid = table._tbl.tblGrid
    for child in list(grid): grid.remove(child)
    for width in widths_dxa:
        col = OxmlElement("w:gridCol"); col.set(qn("w:w"), str(width)); grid.append(col)
    for row in table.rows:
        for idx, cell in enumerate(row.cells):
            tcPr = cell._tc.get_or_add_tcPr()
            tcW = tcPr.find(qn("w:tcW"))
            if tcW is None:
                tcW = OxmlElement("w:tcW"); tcPr.append(tcW)
            tcW.set(qn("w:w"), str(widths_dxa[idx])); tcW.set(qn("w:type"), "dxa")
            set_cell_margins(cell)


def set_run(run, size=11, color=BLACK, bold=None, italic=None, name="Calibri"):
    run.font.name = name
    run._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), name)
    run._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), name)
    run.font.size = Pt(size)
    run.font.color.rgb = RGBColor.from_string(color)
    if bold is not None: run.bold = bold
    if italic is not None: run.italic = italic


def add_field(paragraph, instruction):
    run = paragraph.add_run()
    fldChar1 = OxmlElement("w:fldChar"); fldChar1.set(qn("w:fldCharType"), "begin")
    instrText = OxmlElement("w:instrText"); instrText.set(qn("xml:space"), "preserve"); instrText.text = instruction
    fldChar2 = OxmlElement("w:fldChar"); fldChar2.set(qn("w:fldCharType"), "end")
    run._r.extend([fldChar1, instrText, fldChar2])


def add_bullet(doc, text):
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.line_spacing = 1.208
    p.paragraph_format.left_indent = Inches(0.375)
    p.paragraph_format.first_line_indent = Inches(-0.194)
    p.add_run(text)
    return p


def add_number(doc, text):
    p = doc.add_paragraph(style="List Number")
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.line_spacing = 1.208
    p.paragraph_format.left_indent = Inches(0.375)
    p.paragraph_format.first_line_indent = Inches(-0.194)
    p.add_run(text)
    return p


def add_numbered_list(doc, items):
    """Create a real, independently restarting Word decimal list."""
    numbering = doc.part.numbering_part.element
    abstract_ids = numbering.xpath(
        'w:abstractNum[w:lvl[@w:ilvl="0"]/w:pStyle[@w:val="ListNumber"]]/@w:abstractNumId'
    )
    if not abstract_ids:
        abstract_ids = numbering.xpath('w:abstractNum[w:lvl[@w:ilvl="0"]/w:numFmt[@w:val="decimal"]]/@w:abstractNumId')
    abstract_id = int(abstract_ids[0])
    num = numbering.add_num(abstract_id)
    num.add_lvlOverride(ilvl=0).add_startOverride(1)
    num_id = num.numId

    for item in items:
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(4)
        p.paragraph_format.line_spacing = 1.208
        num_pr = OxmlElement("w:numPr")
        ilvl = OxmlElement("w:ilvl"); ilvl.set(qn("w:val"), "0")
        nid = OxmlElement("w:numId"); nid.set(qn("w:val"), str(num_id))
        num_pr.extend([ilvl, nid]); p._p.get_or_add_pPr().append(num_pr)
        p.add_run(item)


def add_table(doc, headers, rows, widths):
    table = doc.add_table(rows=1, cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table.style = "Table Grid"
    hdr = table.rows[0].cells
    for i, h in enumerate(headers):
        hdr[i].text = h
        set_cell_shading(hdr[i], PALE)
        hdr[i].vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        for run in hdr[i].paragraphs[0].runs: set_run(run, size=9.5, color=NAVY, bold=True)
    header_pr = table.rows[0]._tr.get_or_add_trPr()
    header_repeat = OxmlElement("w:tblHeader"); header_repeat.set(qn("w:val"), "true"); header_pr.append(header_repeat)
    for row in rows:
        cells = table.add_row().cells
        for i, value in enumerate(row):
            cells[i].text = value
            cells[i].vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            for p in cells[i].paragraphs:
                p.paragraph_format.space_after = Pt(2)
                p.paragraph_format.line_spacing = 1.05
                for run in p.runs: set_run(run, size=9.25, color=BLACK)
    set_table_geometry(table, widths)
    for row in table.rows:
        tr_pr = row._tr.get_or_add_trPr()
        cant_split = OxmlElement("w:cantSplit"); cant_split.set(qn("w:val"), "true"); tr_pr.append(cant_split)
    doc.add_paragraph().paragraph_format.space_after = Pt(0)
    return table


def add_callout(doc, label, text):
    table = doc.add_table(rows=1, cols=1)
    table.style = "Table Grid"
    cell = table.cell(0, 0)
    set_cell_shading(cell, PALE)
    p = cell.paragraphs[0]
    r = p.add_run(label + "  "); set_run(r, size=10.5, color=NAVY, bold=True)
    r = p.add_run(text); set_run(r, size=10.5, color=BLACK)
    set_table_geometry(table, [9360])
    doc.add_paragraph().paragraph_format.space_after = Pt(0)


def configure_document(doc):
    sec = doc.sections[0]
    sec.page_width = Inches(8.5); sec.page_height = Inches(11)
    sec.top_margin = sec.right_margin = sec.bottom_margin = sec.left_margin = Inches(1)
    sec.header_distance = sec.footer_distance = Inches(0.492)
    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Calibri"; normal.font.size = Pt(11); normal.font.color.rgb = RGBColor.from_string(BLACK)
    normal._element.rPr.rFonts.set(qn("w:ascii"), "Calibri"); normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    normal.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    normal.paragraph_format.space_before = Pt(0); normal.paragraph_format.space_after = Pt(8); normal.paragraph_format.line_spacing = 1.333
    for name, size, color, before, after in [
        ("Heading 1",16,BLUE,18,10),("Heading 2",13,BLUE,12,6),("Heading 3",12,DARK_BLUE,8,4)
    ]:
        s=styles[name]; s.font.name="Calibri"; s.font.size=Pt(size); s.font.bold=True; s.font.color.rgb=RGBColor.from_string(color)
        s._element.rPr.rFonts.set(qn("w:ascii"),"Calibri"); s._element.rPr.rFonts.set(qn("w:hAnsi"),"Calibri")
        s.paragraph_format.space_before=Pt(before); s.paragraph_format.space_after=Pt(after); s.paragraph_format.keep_with_next=True
    for name in ("List Bullet", "List Number"):
        s=styles[name]; s.font.name="Calibri"; s.font.size=Pt(11)
        s.paragraph_format.left_indent=Inches(0.375); s.paragraph_format.first_line_indent=Inches(-0.194)
        s.paragraph_format.space_after=Pt(4); s.paragraph_format.line_spacing=1.208
    header = sec.header.paragraphs[0]
    header.text = "FLEET OPS COPILOT  |  CONCEPT DESIGN REPORT"
    header.alignment = WD_ALIGN_PARAGRAPH.LEFT
    for r in header.runs: set_run(r, size=8.5, color=MUTED, bold=True)
    footer = sec.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    r=footer.add_run("Fleet Ops Copilot  |  "); set_run(r,size=8.5,color=MUTED)
    add_field(footer,"PAGE")


def build_report():
    doc = Document()
    configure_document(doc)
    props = doc.core_properties
    props.title = "Fleet Ops Copilot - Project Report"
    props.subject = "Concept, objectives, architecture, requirements and delivery roadmap"
    props.author = "Fleet Ops Copilot Project Team"
    props.keywords = "fleet operations, transportation, LLM, agent, schedule, architecture"

    # Editorial cover pattern.
    p=doc.add_paragraph(); p.paragraph_format.space_before=Pt(82); p.paragraph_format.space_after=Pt(16); p.alignment=WD_ALIGN_PARAGRAPH.CENTER
    r=p.add_run("CONCEPT DESIGN REPORT"); set_run(r,size=11,color=GOLD,bold=True)
    p=doc.add_paragraph(); p.paragraph_format.space_after=Pt(8); p.alignment=WD_ALIGN_PARAGRAPH.CENTER
    r=p.add_run("Fleet Ops Copilot"); set_run(r,size=30,color=NAVY,bold=True)
    p=doc.add_paragraph(); p.paragraph_format.space_after=Pt(28); p.alignment=WD_ALIGN_PARAGRAPH.CENTER
    r=p.add_run("An AI agent for live transportation schedules, operational problem solving, and guided action"); set_run(r,size=14,color=DARK_BLUE)
    p=doc.add_paragraph(); p.paragraph_format.space_before=Pt(72); p.paragraph_format.space_after=Pt(6); p.alignment=WD_ALIGN_PARAGRAPH.CENTER
    r=p.add_run("PROJECT VISION"); set_run(r,size=10,color=GOLD,bold=True)
    p=doc.add_paragraph(); p.paragraph_format.left_indent=Inches(.65); p.paragraph_format.right_indent=Inches(.65); p.paragraph_format.space_after=Pt(66); p.alignment=WD_ALIGN_PARAGRAPH.CENTER
    r=p.add_run("Give every transportation operator a trusted conversational interface to current operational data, approved procedures, and safe next actions."); set_run(r,size=13,color=MUTED,italic=True)
    p=doc.add_paragraph(); p.alignment=WD_ALIGN_PARAGRAPH.CENTER
    r=p.add_run("Prepared 13 July 2026  |  Version 1.0"); set_run(r,size=10,color=MUTED,bold=True)
    doc.add_page_break()

    doc.add_heading("Executive Summary", level=1)
    doc.add_paragraph("Fleet Ops Copilot is a proposed transportation operations assistant that combines a large language model (LLM), an agent orchestration layer, and controlled access to live operational systems. It enables dispatchers, planners, drivers, and operations managers to ask questions in natural language, retrieve current schedules, investigate disruptions, and follow approved resolution procedures without navigating multiple disconnected tools.")
    add_callout(doc,"Core design principle", "The copilot should never treat the LLM as the system of record. Live operational services remain authoritative; the agent retrieves, explains, and acts through governed tools.")
    doc.add_paragraph("The proposed design separates the conversational experience from operational tools and data. The agent identifies the user's intent, checks authorization and freshness requirements, selects approved tools, validates returned evidence, and produces an answer with timestamps, confidence, and clear next steps. Any action that changes operations requires explicit approval and a complete audit trail.")

    doc.add_heading("1. Introduction", level=1)
    doc.add_paragraph("Transportation operations depend on accurate schedules, vehicle and driver availability, route conditions, and timely incident handling. This information is commonly distributed across scheduling platforms, fleet systems, ticketing tools, spreadsheets, messages, and operating procedures. Users must know where to look, how to interpret the data, and which action is allowed under policy.")
    doc.add_paragraph("Fleet Ops Copilot provides a single conversational entry point into this environment. A user can ask questions such as 'What is the next available trip?', 'Why is Route 18 delayed?', or 'Which replacement vehicle can cover the 14:30 departure?'. The system gathers evidence from authorized sources, explains the result in operational language, and recommends or performs an approved next step.")

    doc.add_heading("2. Problem Statement", level=1)
    add_bullet(doc,"Operational data is fragmented across systems and may be difficult to reconcile quickly.")
    add_bullet(doc,"Schedule changes and disruptions require rapid decisions, but the latest source and effective time are not always obvious.")
    add_bullet(doc,"Standard operating procedures may be lengthy, inconsistent, or hard to locate during an incident.")
    add_bullet(doc,"Manual handoffs create repeated work, delayed resolution, and incomplete audit records.")
    add_bullet(doc,"A general-purpose chatbot may produce plausible but unsafe answers unless it is grounded in live data and constrained by operational policy.")

    doc.add_heading("3. Project Objectives", level=1)
    objectives = [
        "Provide fast, natural-language access to current transportation schedules, routes, assignments, and service status.",
        "Help users diagnose operational problems using evidence from live systems and approved knowledge sources.",
        "Recommend safe next actions that follow organizational policy and explain the reasoning behind each recommendation.",
        "Automate low-risk tasks through governed tools while requiring human approval for changes with operational impact.",
        "Reduce time spent switching between systems and improve consistency of incident handling.",
        "Create a searchable audit trail for questions, sources, decisions, approvals, actions, and outcomes.",
        "Measure answer quality, operational value, safety, latency, and user adoption through defined service metrics.",
    ]
    add_numbered_list(doc, objectives)

    doc.add_heading("4. Scope", level=1)
    doc.add_heading("4.1 In Scope", level=2)
    for item in ["Schedule and route lookup using live data.","Vehicle and driver availability queries based on user role.","Disruption triage and guided incident resolution.","Retrieval of policies, standard operating procedures, and route knowledge.","Recommendations, summaries, notifications, and approved ticket or schedule actions.","Role-based access, evidence display, human approval, logging, and monitoring."]:
        add_bullet(doc,item)
    doc.add_heading("4.2 Initial Exclusions", level=2)
    for item in ["Autonomous dispatch decisions without human oversight.","Direct control of vehicle safety systems.","Payroll, disciplinary, or other sensitive HR decisions.","Prediction claims that cannot be validated against operational data.","Replacement of authoritative scheduling, fleet, or incident-management systems."]:
        add_bullet(doc,item)

    heading = doc.add_heading("5. Users and Key Use Cases", level=1)
    heading.paragraph_format.page_break_before = True
    add_table(doc,["User","Primary need","Example copilot request"],[
        ("Dispatcher","Current operational picture","Show delayed departures in the next two hours and their assigned vehicles."),
        ("Planner","Schedule and capacity analysis","Which trips on Route 12 are likely to exceed available capacity?"),
        ("Driver / field team","Clear, role-appropriate guidance","My assigned vehicle is unavailable. What is the approved next step?"),
        ("Operations manager","Exceptions, trends, and decisions","Summarize today's major disruptions, actions taken, and unresolved risks."),
        ("Support analyst","Consistent incident resolution","Create an incident from this conversation and attach the evidence."),
    ],[1500,2860,5000])

    doc.add_heading("6. Functional Requirements", level=1)
    add_table(doc,["ID","Requirement","Priority"],[
        ("FR-01","Authenticate the user and determine role, organization, and permitted data scope.","Must"),
        ("FR-02","Retrieve current schedules and show source system plus effective timestamp.","Must"),
        ("FR-03","Classify intent and select only tools approved for that intent and role.","Must"),
        ("FR-04","Retrieve relevant SOP and policy content with traceable references.","Must"),
        ("FR-05","Detect stale, missing, or conflicting data and communicate uncertainty.","Must"),
        ("FR-06","Require explicit approval before executing an operational change.","Must"),
        ("FR-07","Record prompts, sources, tool calls, approvals, actions, and results.","Must"),
        ("FR-08","Escalate to a human when confidence or tool availability falls below policy.","Must"),
        ("FR-09","Support multilingual or voice channels where operationally justified.","Could"),
    ],[900,7100,1360])

    doc.add_heading("7. Proposed System Architecture", level=1)
    doc.add_paragraph("The architecture follows a layered model that keeps the LLM isolated from direct database access. Every live query or operational action is performed through a typed, authenticated tool with validation, timeouts, and auditable results.")
    architecture_shape = doc.add_picture(str(ARCH_PNG), width=Inches(6.5))
    architecture_shape._inline.docPr.set(
        "descr",
        "Fleet Ops Copilot layered architecture connecting users to the copilot experience, agent orchestration, operational tools, data and knowledge, and platform controls.",
    )
    architecture_shape._inline.docPr.set("title", "Fleet Ops Copilot reference architecture")
    p=doc.add_paragraph("Figure 1. Fleet Ops Copilot reference architecture. The editable source is provided in the companion Draw.io file.")
    p.alignment=WD_ALIGN_PARAGRAPH.CENTER; p.paragraph_format.space_before=Pt(4); p.paragraph_format.space_after=Pt(8)
    for r in p.runs: set_run(r,size=9,color=MUTED,italic=True)
    add_table(doc,["Layer","Responsibility"],[
        ("Channels and experience","Conversation, identity context, evidence presentation, approvals, and accessible user interaction."),
        ("Agent orchestration","Intent detection, planning, tool selection, policy enforcement, grounding, synthesis, and escalation."),
        ("Operational tools","Stable interfaces to schedule, route, fleet, driver, incident, traffic, weather, and notification services."),
        ("Data and knowledge","Authoritative operational stores, indexed SOPs, route knowledge, session state, and audit records."),
        ("Platform controls","API gateway, role-based access, secrets, encryption, monitoring, cost controls, retention, and compliance."),
    ],[2100,7260])

    doc.add_heading("8. Agent Workflow", level=1)
    workflow = [
        "Receive the request and capture conversation context, channel, locale, and requested outcome.",
        "Authenticate the user, apply role and organization scope, and identify data restrictions.",
        "Classify the intent, determine the required data freshness, and identify safety or policy constraints.",
        "Create a bounded plan and select only approved read or action tools.",
        "Query live operational systems and retrieve relevant SOP or policy passages.",
        "Validate timestamps, reconcile conflicting sources, calculate confidence, and detect missing evidence.",
        "Compose a concise answer with operational impact, evidence, assumptions, and recommended next action.",
        "If the next action changes operations, request explicit approval; then execute through the tool and record the result.",
        "Confirm the outcome or escalate to a human with the conversation, evidence, and failed checks attached.",
    ]
    add_numbered_list(doc, workflow)
    add_callout(doc,"Fallback behavior", "When live data cannot be reached, the copilot must say so clearly, avoid presenting cached data as current, and provide a safe escalation path.")

    doc.add_heading("9. Data and Integration Design", level=1)
    doc.add_paragraph("A canonical operational model should normalize identifiers and timestamps across connected systems. At minimum, the model should represent trip, route, stop, service calendar, vehicle, driver, assignment, incident, disruption, action, source, and user. Each retrieved fact should carry provenance, observed time, effective time, and freshness status.")
    add_table(doc,["Integration","Typical data","Access pattern"],[
        ("Schedule service","Trips, stops, planned times, revisions","Read-heavy; event updates"),
        ("Fleet / telematics","Vehicle status, position, availability","Read; near-real-time stream"),
        ("Workforce / roster","Driver assignment and eligibility","Restricted read"),
        ("Incident platform","Cases, status, notes, ownership","Read and approved write"),
        ("Maps / external feeds","Traffic, weather, travel time","Read with provider timestamp"),
        ("Knowledge repository","SOPs, policies, route notes","Indexed retrieval with versioning"),
    ],[2200,3900,3260])

    doc.add_heading("10. Safety, Security, and Governance", level=1)
    for item in [
        "Least privilege: each tool enforces user, role, tenant, route, and action scope independently of the model.",
        "Grounding: operational claims must be linked to retrieved facts; schedule answers display source and timestamp.",
        "Approval: changes to schedules, assignments, incidents, or outbound notifications require policy-based confirmation.",
        "Data protection: encrypt data in transit and at rest, keep secrets outside prompts, and minimize personal information.",
        "Prompt-injection defense: treat external text as untrusted data, isolate tool instructions, and validate every parameter.",
        "Auditability: record tool inputs and outputs, policy decisions, approvals, model version, and final response.",
        "Retention: define separate retention periods for conversations, operational evidence, analytics, and sensitive data.",
        "Human escalation: route safety-critical, ambiguous, or policy-exception decisions to an accountable operator.",
    ]: add_bullet(doc,item)

    doc.add_heading("11. Non-Functional Requirements", level=1)
    add_table(doc,["Area","Initial target"],[
        ("Availability","99.5% for the copilot channel; graceful degradation when a source system is unavailable."),
        ("Latency","Typical read-only answer within 5 seconds; progressive status for longer multi-tool investigations."),
        ("Freshness","Schedule responses identify observed and effective time; stale thresholds are defined per source."),
        ("Scalability","Stateless orchestration workers, queued long-running tasks, and rate-limited external integrations."),
        ("Reliability","Idempotent action tools, retries only for safe operations, and reconciliation after partial failure."),
        ("Accessibility","Keyboard navigation, screen-reader labels, readable evidence, and non-color-only status indicators."),
        ("Maintainability","Versioned prompts, schemas, tool contracts, policies, evaluation sets, and deployment artifacts."),
    ],[2200,7160])

    doc.add_heading("12. Delivery Roadmap", level=1)
    add_table(doc,["Phase","Focus","Exit criteria"],[
        ("1. Discovery","Systems, users, decisions, data classification, and baseline metrics.","Approved use cases, source inventory, risk assessment."),
        ("2. Read-only MVP","Schedule lookup, SOP retrieval, evidence, access control, and audit.","Pilot users achieve target accuracy and latency with no ungrounded operational claims."),
        ("3. Guided resolution","Incident diagnosis, recommendations, confidence, and escalation.","Defined scenarios pass offline evaluation and supervised field trials."),
        ("4. Approved actions","Ticket creation, notifications, and selected schedule workflows.","Approval, idempotency, rollback, and audit controls validated."),
        ("5. Scale and optimize","More routes, sources, channels, analytics, and cost optimization.","Service objectives met under production load and governance review."),
    ],[1300,3860,4200])

    doc.add_heading("13. Success Metrics", level=1)
    for item in [
        "Schedule-answer accuracy and percentage of answers supported by current authoritative data.",
        "Median and 95th-percentile response time for common operational questions.",
        "Mean time to understand and resolve selected disruption types.",
        "Tool-call success, action confirmation, and partial-failure recovery rates.",
        "Escalation precision: unsafe or ambiguous cases escalated without excessive false alarms.",
        "User adoption, repeated usage, task completion, and operator satisfaction.",
        "Safety indicators: unauthorized actions, stale-data incidents, unsupported claims, and audit completeness.",
        "Cost per completed operational task compared with the baseline workflow.",
    ]: add_bullet(doc,item)

    doc.add_heading("14. Risks and Mitigations", level=1)
    add_table(doc,["Risk","Mitigation"],[
        ("Hallucinated or unsupported answer","Require tool-derived evidence, confidence checks, response contracts, and evaluation against known scenarios."),
        ("Stale or conflicting schedule data","Display effective time, apply source precedence, detect conflicts, and block unsafe conclusions."),
        ("Unauthorized information or action","Enforce identity and role at the gateway and again inside every tool."),
        ("Prompt injection through documents or feeds","Separate data from instructions, sanitize content, use allow-listed tools, and validate parameters."),
        ("Partial operational change","Use idempotency keys, transaction boundaries where possible, reconciliation, and visible action receipts."),
        ("Over-reliance on automation","Keep accountable humans in change workflows and train users on limitations and escalation."),
        ("Low trust or adoption","Show sources and reasoning, start with high-value read-only use cases, and incorporate operator feedback."),
    ],[3000,6360])

    doc.add_heading("15. Assumptions and Open Decisions", level=1)
    doc.add_paragraph("This report is a concept baseline because the current repository does not yet contain implementation requirements or source-system details. The following decisions should be confirmed during discovery:")
    for item in [
        "Transportation mode and operating region, including any sector-specific regulation.",
        "Authoritative scheduling, fleet, roster, telematics, incident, and knowledge systems.",
        "User roles, tenant boundaries, personal data classification, and retention rules.",
        "Required channels, languages, response-time objectives, and peak concurrency.",
        "Which actions may be automated, which require approval, and who is accountable for exceptions.",
        "Preferred LLM provider, deployment environment, integration standards, and operational budget.",
    ]: add_bullet(doc,item)

    heading = doc.add_heading("16. Conclusion", level=1)
    heading.paragraph_format.page_break_before = True
    doc.add_paragraph("Fleet Ops Copilot can reduce operational friction by making current schedules, fleet information, and approved procedures accessible through one trusted conversational interface. The value does not come from the LLM alone; it comes from a governed agent that connects the user to authoritative systems, exposes evidence and freshness, respects role boundaries, and keeps humans accountable for consequential changes.")
    doc.add_paragraph("A read-only MVP focused on schedule access and SOP retrieval is the recommended starting point. Once accuracy, latency, access control, and operator trust are demonstrated, the platform can progress to guided incident resolution and carefully selected approved actions.")
    add_callout(doc,"Recommended next step", "Run a short discovery phase with dispatchers and operations managers, select three high-volume scenarios, and define the authoritative data source, safety policy, and measurable baseline for each scenario.")
    doc.add_heading("Immediate next actions", level=2)
    for item in [
        "Nominate a product owner and an accountable transportation operations lead.",
        "Choose three high-volume, low-risk pilot scenarios and document their current workflow.",
        "Confirm the authoritative system, identifier, timestamp, and owner for every required data element.",
        "Create a representative evaluation set containing normal, ambiguous, stale-data, and failure cases.",
        "Approve read-only MVP boundaries, escalation rules, service targets, and the production-readiness gate.",
    ]: add_bullet(doc, item)

    doc.save(DOCX_PATH)


if __name__ == "__main__":
    draw_architecture_png()
    build_drawio()
    build_report()
    print(DOCX_PATH)
    print(DRAWIO_PATH)
